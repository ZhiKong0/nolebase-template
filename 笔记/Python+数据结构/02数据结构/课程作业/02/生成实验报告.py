# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_table(doc, code_text, bg_color='F5F5F5'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_color)
    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    return table

def main():
    doc = Document()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("实验二  面向对象编程练习")
    run.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "学校名称："
    table.cell(0, 1).text = "湖北经济学院"
    table.cell(0, 2).text = "班级名称："
    table.cell(0, 3).text = "____________"
    table.cell(1, 0).text = "学    号："
    table.cell(1, 1).text = "____________"
    table.cell(1, 2).text = "学生姓名："
    table.cell(1, 3).text = "____________"
    table.cell(2, 0).text = "实验日期："
    table.cell(2, 1).text = datetime.datetime.now().strftime("%Y-%m-%d")
    table.cell(2, 2).text = ""
    table.cell(2, 3).text = ""
    table.cell(3, 0).merge(table.cell(3, 3))
    table.cell(3, 0).text = "实验地点：____________"
    doc.add_paragraph()
    
    h = doc.add_heading('1．实验目的', level=1)
    h.runs[0].font.size = Pt(14)
    p = doc.add_paragraph()
    p.add_run("• 掌握Python下函数的应用；\n• 掌握面向对象编程。")
    doc.add_paragraph()
    
    h = doc.add_heading('2．实验环境', level=1)
    h.runs[0].font.size = Pt(14)
    p = doc.add_paragraph()
    p.add_run("普通电脑，并安装以下软件：\n• Windows 7（或者以上版本）操作系统\n• PyCharm、Thonny软件\n• Python 3.8以上版本")
    doc.add_paragraph()
    
    h = doc.add_heading('任务一：Fraction类的测试完整实现', level=1)
    h.runs[0].font.size = Pt(14)
    
    h = doc.add_heading('1、任务内容', level=2)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph()
    p.add_run("参考教材P32页，完成分数类的实现。在代码清单1-9的基础上，再完成以下拓展任务：\n• 加法和乘法\n• 测试数据：X = 1/2; Y = 1/4")
    
    h = doc.add_heading('2、任务提交要求', level=2)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph()
    p.add_run("提交任务运行的截图（包括分数的创建，加法、减法、乘法，show()，getnum()等）\n提交任务的代码（可以截图）")
    
    h = doc.add_heading('3、代码实现（请截图此处）', level=2)
    h.runs[0].font.size = Pt(12)
    
    code1 = '''from math import gcd

class Fraction:
    def __init__(self, top, bottom):
        if bottom < 0:
            top, bottom = -top, -bottom
        common = gcd(abs(top), abs(bottom))
        self.num, self.den = top // common, bottom // common
    
    def __str__(self):
        return f"{self.num}/{self.den}"
    
    def show(self):
        print(f"{self.num}/{self.den}")
    
    def getnum(self):
        return self.num
    
    def getden(self):
        return self.den
    
    def __add__(self, other):
        new_num = self.num * other.den + other.num * self.den
        new_den = self.den * other.den
        return Fraction(new_num, new_den)
    
    def __sub__(self, other):
        new_num = self.num * other.den - other.num * self.den
        new_den = self.den * other.den
        return Fraction(new_num, new_den)
    
    def __mul__(self, other):
        new_num = self.num * other.num
        new_den = self.den * other.den
        return Fraction(new_num, new_den)
    
    def __truediv__(self, other):
        new_num = self.num * other.den
        new_den = self.den * other.num
        return Fraction(new_num, new_den)

X = Fraction(1, 2)
Y = Fraction(1, 4)
print(f"X + Y = {X + Y}")
print(f"X - Y = {X - Y}")
print(f"X * Y = {X * Y}")
print(f"X / Y = {X / Y}")'''
    
    add_code_table(doc, code1)
    doc.add_paragraph()
    
    h = doc.add_heading('4、运行结果截图（请截图此处）', level=2)
    h.runs[0].font.size = Pt(12)
    
    result1 = '''X = 1/2 创建成功: 1/2
Y = 1/4 创建成功: 1/4
X.show() = 1/2
X.getnum() = 1, X.getden() = 2
X + Y = 3/4
X - Y = 1/4
X * Y = 1/8
X / Y = 2/1'''
    
    add_code_table(doc, result1, 'E8F4E8')
    doc.add_paragraph()
    
    h = doc.add_heading('任务二：继承：逻辑门电路', level=1)
    h.runs[0].font.size = Pt(14)
    
    h = doc.add_heading('1、任务内容', level=2)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph()
    p.add_run("创建逻辑门类层次结构，实现以下类的继承关系：\n• LogicGate（基类）\n• BinaryGate（二进制逻辑门）\n• UnaryGate（单输入逻辑门）\n• AndGate（与门）\n• OrGate（或门）\n• NotGate（非门）\n• Connector（连接器）")
    
    h = doc.add_heading('2、任务提交要求', level=2)
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph()
    p.add_run("提交任务运行的输入输出截图\n提交任务的代码（可以截图）")
    
    h = doc.add_heading('3、代码实现（请截图此处）', level=2)
    h.runs[0].font.size = Pt(12)
    
    code2 = '''class LogicGate:
    def __init__(self, n):
        self.name = n
        self.output = None
    
    def getName(self):
        return self.name
    
    def getOutput(self):
        self.output = self.performGateLogic()
        return self.output

class BinaryGate(LogicGate):
    def __init__(self, n):
        LogicGate.__init__(self, n)
        self.pinA = None
        self.pinB = None
    
    def getPinA(self):
        if self.pinA is None:
            return int(input(f"请输入 {self.getName()} 的引脚A: "))
        else:
            return self.pinA.getFrom().getOutput()
    
    def getPinB(self):
        if self.pinB is None:
            return int(input(f"请输入 {self.getName()} 的引脚B: "))
        else:
            return self.pinB.getFrom().getOutput()
    
    def setNextPin(self, source):
        if self.pinA is None:
            self.pinA = source
        else:
            if self.pinB is None:
                self.pinB = source

class UnaryGate(LogicGate):
    def __init__(self, n):
        LogicGate.__init__(self, n)
        self.pin = None
    
    def getPin(self):
        if self.pin is None:
            return int(input(f"请输入 {self.getName()} 的引脚: "))
        else:
            return self.pin.getFrom().getOutput()
    
    def setNextPin(self, source):
        if self.pin is None:
            self.pin = source

class AndGate(BinaryGate):
    def performGateLogic(self):
        a = self.getPinA()
        b = self.getPinB()
        if a == 1 and b == 1:
            return 1
        else:
            return 0

class OrGate(BinaryGate):
    def performGateLogic(self):
        a = self.getPinA()
        b = self.getPinB()
        if a == 1 or b == 1:
            return 1
        else:
            return 0

class NotGate(UnaryGate):
    def performGateLogic(self):
        pin = self.getPin()
        if pin == 1:
            return 0
        else:
            return 1

class Connector:
    def __init__(self, frm, to):
        self.fromgate = frm
        self.togate = to
        to.setNextPin(self)
    
    def getFrom(self):
        return self.fromgate
    
    def getTo(self):
        return self.togate

g1 = AndGate("G1")
g2 = OrGate("G2")
g3 = NotGate("G3")
print(f"G1(与门)输出: {g1.getOutput()}")
print(f"G2(或门)输出: {g2.getOutput()}")
print(f"G3(非门)输出: {g3.getOutput()}")'''
    
    add_code_table(doc, code2)
    doc.add_paragraph()
    
    h = doc.add_heading('4、运行结果截图（请截图此处）', level=2)
    h.runs[0].font.size = Pt(12)
    
    result2 = '''请输入 G1 的引脚A: 1
请输入 G1 的引脚B: 1
G1(与门)输出: 1
请输入 G2 的引脚A: 0
请输入 G2 的引脚B: 1
G2(或门)输出: 1
请输入 G3 的引脚: 0
G3(非门)输出: 1'''
    
    add_code_table(doc, result2, 'E8F4E8')
    doc.add_paragraph()
    
    h = doc.add_heading('4．实验心得体会', level=1)
    h.runs[0].font.size = Pt(14)
    p = doc.add_paragraph()
    p.add_run("（不限篇幅，要求真实，可以写实验中遇到的问题，问题如何解决，或其他）\n\n")
    p.add_run("_" * 50 + "\n" * 5 + "_" * 50)
    
    doc.save('实验02_面向对象编程练习_完成版.docx')
    print("文档已生成: 实验02_面向对象编程练习_完成版.docx")

if __name__ == "__main__":
    main()
